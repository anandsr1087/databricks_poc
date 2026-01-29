#!/usr/bin/env python3
"""
Convert Markdown Use Case Document to Word format
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_page_number(section):
    """Add page numbers to the document"""
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page "
    from docx.oxml import OxmlElement
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('w:val'), 'PAGE')
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def parse_table_row(line):
    """Parse a markdown table row and return list of cell contents"""
    cells = [c.strip() for c in line.split('|')[1:-1]]
    # Clean up markdown formatting in cells
    cleaned_cells = []
    for cell in cells:
        # Remove bold markers for clean text
        cell = re.sub(r'\*\*(.+?)\*\*', r'\1', cell)
        cell = re.sub(r'\*(.+?)\*', r'\1', cell)
        cell = cell.replace('\\', '')
        cleaned_cells.append(cell)
    return cleaned_cells

def is_table_separator(line):
    """Check if line is a markdown table separator"""
    return bool(re.match(r'^\|[\s\:]+\|[\s\:]+\|', line))

def main():
    # Read the markdown file
    input_file = '/Users/anand.kumar/.Trash/databricks-poc-/docs/USE_CASE_DOCUMENT.md'
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create a new Document
    doc = Document()

    # Set up default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title page
    title = doc.add_heading('USE CASE DOCUMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Databricks Retail Analytics Proof of Concept', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Document Information Section
    doc.add_heading('Document Information', 2)

    info_data = [
        ['Project Name', 'Databricks - Retail Analytics Platform'],
        ['Document Version', '1.0'],
        ['Last Updated', '2026-01-29'],
        ['Status', 'Draft'],
        ['Confidentiality', 'Internal']
    ]

    for key, value in info_data:
        p = doc.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(value)

    doc.add_page_break()

    # Process the content
    lines = content.split('\n')
    i = 0
    in_code_block = False
    skip_until_next_heading = False

    while i < len(lines):
        line = lines[i]

        # Skip already processed sections
        if 'USE CASE DOCUMENT' in line and line.startswith('#'):
            i += 1
            continue

        if '**Document Information**' in line:
            skip_until_next_heading = True
            i += 1
            continue

        # Reset skip flag at next heading
        if skip_until_next_heading and line.startswith('#'):
            skip_until_next_heading = False

        if skip_until_next_heading:
            i += 1
            continue

        # Code blocks (ASCII diagrams)
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                p = doc.add_paragraph()
                run = p.add_run('▌ Code Block')
                run.font.color.rgb = RGBColor(100, 100, 100)
                run.font.size = Pt(9)
            i += 1
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('─' * 80)
            i += 1
            continue

        # Tables
        if line.startswith('|') and not is_table_separator(line):
            # Collect all table rows
            table_lines = []
            start_i = i
            while i < len(lines) and lines[i].startswith('|'):
                if not is_table_separator(lines[i]):
                    table_lines.append(lines[i])
                i += 1

            # Parse and create table
            if table_lines:
                rows = [parse_table_row(tl) for tl in table_lines]
                if rows and all(len(r) == len(rows[0]) for r in rows):
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    table.autofit = True

                    for ri, row_data in enumerate(rows):
                        row = table.rows[ri]
                        for ci in range(min(len(row_data), len(row.cells))):
                            cell = row.cells[ci]
                            cell.text = row_data[ci]
                            # Bold header rows
                            if ri == 0 or '**' in table_lines[ri]:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.size = Pt(10)
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

        # Headers
        if line.startswith('#'):
            level_match = re.match(r'^(#+)\s+(.+)', line)
            if level_match:
                level = len(level_match.group(1))
                text = level_match.group(2).strip()
                doc.add_heading(text, min(level, 3))
            i += 1
            continue

        # Bullet lists
        bullet_match = re.match(r'^\s*[\*\-\+]\s+(.+)', line)
        if bullet_match:
            p = doc.add_paragraph(bullet_match.group(1), style='List Bullet')
            i += 1
            continue

        # Numbered lists
        num_match = re.match(r'^\s*\d+\.\s+(.+)', line)
        if num_match:
            p = doc.add_paragraph(num_match.group(1), style='List Number')
            i += 1
            continue

        # Regular paragraphs
        if line.strip():
            # Handle bold text
            if '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'\*\*(.+?)\*\*', line)
                for idx, part in enumerate(parts):
                    run = p.add_run(part.strip())
                    run.bold = (idx % 2 == 1)
            else:
                doc.add_paragraph(line.strip())

        i += 1

    # Configure page settings
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_page_number(section)

    # Save the document
    output_path = '/Users/anand.kumar/.Trash/databricks-poc-/docs/USE_CASE_DOCUMENT.docx'
    doc.save(output_path)
    print(f"✓ Word document saved to: {output_path}")
    print(f"  Location: {output_path}")

if __name__ == '__main__':
    main()
